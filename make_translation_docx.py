from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"D:\python_project\译文.txt")
TARGET = Path(r"D:\python_project\译文.docx")


def set_run_font(run, size=11, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(document, text, level):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10 if level else 0)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 0 else 13, bold=True)


def add_body_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def main():
    text = SOURCE.read_text(encoding="utf-8-sig")
    blocks = [block.strip() for block in text.splitlines() if block.strip()]

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)

    major_headings = {
        "摘要",
        "引言",
        "结果",
        "方法",
        "讨论",
        "相关工作",
        "伦理考虑",
        "招募说明",
        "数据可用性",
        "代码可用性",
        "致谢",
        "作者贡献",
        "利益冲突",
        "附加信息",
    }

    sub_headings = {
        "数据集、指标与设置",
        "疾病报告生成结果",
        "疾病分类结果",
        "框架",
        "MultiMedCLIP",
        "MultiMedLM",
        "训练设置",
    }

    for i, block in enumerate(blocks):
        if i == 0:
            add_heading(document, block, 0)
        elif block in major_headings or block in sub_headings:
            add_heading(document, block, 1)
        else:
            add_body_paragraph(document, block)

    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
